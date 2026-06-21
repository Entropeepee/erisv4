"""
eris/knowledge/author.py
========================
Professional document authoring (Eris writes for you). Instead of one shot, Eris
works like a writer: she PLANS an outline, writes each section in full, then
AUDITS and revises the whole draft before it's offered to you. Output to
Markdown, plain text, Word (.docx) and PDF — no length cap (each section is its
own generation, so the document grows section by section).

Reuses her language model (the mediator). Heavy/optional exporters degrade
gracefully: .md/.txt always work; .docx needs python-docx; .pdf needs reportlab.
"""
from __future__ import annotations

import json
import os
import re
import time
from typing import Any, Dict, List, Optional

from eris.interface.mediator import run_blocking


def _slug(s: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "-", (s or "document").strip()).strip("-").lower()
    return (s or "document")[:60]


def _strip_md(md: str) -> str:
    """Markdown → readable plain text (for the .txt export)."""
    t = md or ""
    t = re.sub(r"`{1,3}", "", t)
    t = re.sub(r"^\s{0,3}#{1,6}\s*", "", t, flags=re.M)
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"\*([^*]+)\*", r"\1", t)
    t = re.sub(r"^\s*[-*]\s+", "  • ", t, flags=re.M)
    return t


class DocumentAuthor:
    def __init__(self, mediator, *, out_dir: str):
        self.mediator = mediator
        self.out_dir = out_dir
        os.makedirs(out_dir, exist_ok=True)
        self.progress: Dict[str, Any] = {
            "running": False, "phase": "idle", "title": "",
            "step": 0, "total": 0, "updated": 0.0}

    # ── LLM helper ───────────────────────────────────────────────────────
    def _gen(self, prompt: str, system: str = "") -> str:
        if not self.mediator:
            return ""
        try:
            resp = run_blocking(
                self.mediator.generate(prompt=prompt, system=system), timeout=300)
            return (getattr(resp, "text", "") or "").strip()
        except Exception:
            return ""

    def _set(self, **kw) -> None:
        self.progress.update(updated=time.time(), **kw)

    # ── 1. PLAN ──────────────────────────────────────────────────────────
    def plan(self, brief: str) -> Dict[str, Any]:
        """Outline = her chain of thought: a title + ordered sections."""
        system = ("You are a professional writer and editor planning a document "
                  "before writing it. Think about structure, completeness and flow.")
        prompt = (f"Plan a thorough, professional document for this request:\n\n"
                  f"{brief}\n\n"
                  "Return ONLY a JSON object of the form:\n"
                  '{"title": "...", "sections": [{"heading": "...", '
                  '"intent": "one line on what this section covers"}]}\n'
                  "Use 4-12 well-ordered sections that fully cover the request.")
        raw = self._gen(prompt, system)
        title, sections = self._parse_plan(raw, brief)
        return {"title": title, "sections": sections}

    def _parse_plan(self, raw: str, brief: str):
        title = (brief.strip().split("\n")[0][:90] or "Document")
        sections: List[Dict[str, str]] = []
        m = re.search(r"\{.*\}", raw or "", flags=re.S)
        if m:
            try:
                data = json.loads(m.group(0))
                title = (data.get("title") or title)[:120]
                for s in data.get("sections", []):
                    if isinstance(s, dict) and s.get("heading"):
                        sections.append({"heading": str(s["heading"])[:120],
                                         "intent": str(s.get("intent", ""))[:200]})
            except Exception:
                pass
        if not sections:  # fallback: any markdown/bullet headings in the reply
            for line in (raw or "").splitlines():
                h = re.sub(r"^\s*(#{1,6}|[-*\d.]+)\s*", "", line).strip()
                if 3 < len(h) < 120:
                    sections.append({"heading": h, "intent": ""})
            sections = sections[:12]
        if not sections:
            sections = [{"heading": "Overview", "intent": ""},
                        {"heading": "Details", "intent": ""},
                        {"heading": "Conclusion", "intent": ""}]
        return title, sections

    # ── 2. WRITE each section ────────────────────────────────────────────
    def write_section(self, brief, title, outline_text, heading, intent,
                      prior_summary) -> str:
        system = ("You are a professional writer producing ONE section of a larger "
                  "document. Write substantive, accurate, well-structured prose. Do "
                  "not summarize or truncate; do not write other sections.")
        prompt = (f"Document title: {title}\n\nOriginal request:\n{brief}\n\n"
                  f"Full outline:\n{outline_text}\n\n"
                  f"Sections written so far (brief summaries):\n"
                  f"{prior_summary or '(none yet)'}\n\n"
                  f"Now write the section \"{heading}\" in full"
                  + (f" (intent: {intent})" if intent else "") + ".\n"
                  f"Start with a Markdown header `## {heading}`, then write thorough, "
                  "professional Markdown prose. Be complete.")
        return self._gen(prompt, system)

    # ── 3. AUDIT / revise ────────────────────────────────────────────────
    def audit(self, brief, title, full_md) -> str:
        # Skip the global pass for very long docs (would exceed context); the
        # outline + running summaries already keep long docs coherent.
        if len(full_md) > 16000:
            return full_md
        system = ("You are a meticulous editor. Improve flow and coherence, remove "
                  "redundancy, and ensure the request is fully met — keep all "
                  "substance and length.")
        prompt = (f"Final-edit this document titled '{title}' for the request:\n\n"
                  f"{brief}\n\n--- DRAFT ---\n{full_md}\n--- END ---\n\n"
                  "Return ONLY the final, polished Markdown document.")
        revised = self._gen(prompt, system)
        return revised if len(revised) > 0.5 * len(full_md) else full_md

    # ── compose: the whole pipeline ──────────────────────────────────────
    def compose(self, brief: str, *, formats=("md",), do_audit: bool = True,
                plan: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        formats = [f.lower() for f in formats] or ["md"]
        self._set(running=True, phase="planning", step=0, total=0, title="")
        plan = plan or self.plan(brief)
        title = plan["title"]
        sections = plan["sections"]
        outline_text = "\n".join(f"{i+1}. {s['heading']}"
                                 + (f" — {s['intent']}" if s.get("intent") else "")
                                 for i, s in enumerate(sections))
        self._set(phase="writing", title=title, total=len(sections) + 1, step=0)

        body, prior = [], ""
        for i, s in enumerate(sections):
            text = self.write_section(brief, title, outline_text,
                                      s["heading"], s.get("intent", ""), prior)
            if text:
                body.append(text if text.lstrip().startswith("#")
                            else f"## {s['heading']}\n\n{text}")
                prior += f"\n- {s['heading']}: {text[:250]}"
            self._set(step=i + 1)

        full_md = f"# {title}\n\n" + "\n\n".join(body)
        if do_audit:
            self._set(phase="auditing", step=len(sections))
            full_md = self.audit(brief, title, full_md)
        self._set(step=len(sections) + 1, phase="exporting")

        files = self._export(title, full_md, formats)
        self._set(running=False, phase="done")
        return {"title": title, "outline": sections, "markdown": full_md,
                "files": files}

    # ── exporters ────────────────────────────────────────────────────────
    def _export(self, title: str, md: str, formats) -> List[Dict[str, str]]:
        base = f"{_slug(title)}-{time.strftime('%Y%m%d-%H%M%S')}"
        out = []
        for fmt in formats:
            try:
                if fmt == "md":
                    path = self._write_text(base + ".md", md)
                elif fmt == "txt":
                    path = self._write_text(base + ".txt", _strip_md(md))
                elif fmt == "docx":
                    path = self._to_docx(base + ".docx", title, md)
                elif fmt == "pdf":
                    path = self._to_pdf(base + ".pdf", title, md)
                else:
                    continue
                out.append({"name": os.path.basename(path), "format": fmt,
                            "path": path})
            except Exception as e:
                out.append({"name": base + "." + fmt, "format": fmt, "error": str(e)})
        return out

    def _write_text(self, name: str, text: str) -> str:
        path = os.path.join(self.out_dir, name)
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
        return path

    def _to_docx(self, name: str, title: str, md: str) -> str:
        import docx  # python-docx
        doc = docx.Document()
        doc.add_heading(title, 0)
        for line in md.splitlines():
            s = line.rstrip()
            if not s.strip():
                continue
            h = re.match(r"^(#{1,6})\s+(.*)", s)
            if h:
                lvl = min(len(h.group(1)), 4)
                if h.group(2).strip() != title:
                    doc.add_heading(h.group(2).strip(), lvl)
            elif re.match(r"^\s*[-*]\s+", s):
                doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", s), style="List Bullet")
            else:
                doc.add_paragraph(re.sub(r"\*\*([^*]+)\*\*", r"\1", s))
        path = os.path.join(self.out_dir, name)
        doc.save(path)
        return path

    def _to_pdf(self, name: str, title: str, md: str) -> str:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        styles = getSampleStyleSheet()
        path = os.path.join(self.out_dir, name)
        doc = SimpleDocTemplate(path, pagesize=letter,
                                title=title, leftMargin=0.9 * inch,
                                rightMargin=0.9 * inch)
        flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

        def esc(x):
            return (x.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        for line in md.splitlines():
            s = line.rstrip()
            if not s.strip():
                flow.append(Spacer(1, 6)); continue
            h = re.match(r"^(#{1,6})\s+(.*)", s)
            if h:
                if h.group(2).strip() == title:
                    continue
                lvl = min(len(h.group(1)), 3)
                flow.append(Spacer(1, 8))
                flow.append(Paragraph(esc(h.group(2).strip()), styles[f"Heading{lvl}"]))
            else:
                body = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", esc(s))
                body = re.sub(r"^\s*[-*]\s+", "• ", body)
                flow.append(Paragraph(body, styles["BodyText"]))
        doc.build(flow)
        return path

    def list_documents(self) -> List[Dict[str, Any]]:
        rows = []
        for fn in sorted(os.listdir(self.out_dir), reverse=True):
            p = os.path.join(self.out_dir, fn)
            if os.path.isfile(p):
                rows.append({"name": fn, "size": os.path.getsize(p),
                             "modified": os.path.getmtime(p)})
        return rows
